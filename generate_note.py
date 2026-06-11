"""Генератор пояснительной записки к дипломному проекту DigitalStore."""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime


def set_cell_shading(cell, color):
    """Установить цвет фона ячейки."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def set_cell_borders(cell, top="000000", bottom="000000", left="000000", right="000000"):
    """Установить границы ячейки."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), val)
        tcBorders.append(border)
    tcPr.append(tcBorders)


def add_formatted_paragraph(doc, text, bold=False, size=12, alignment=WD_ALIGN_PARAGRAPH.LEFT,
                            space_after=6, space_before=0, font_name='Times New Roman',
                            color=None, italic=False):
    """Добавить форматированный абзац."""
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = RGBColor(*color)
    return p


def add_heading_styled(doc, text, level=1):
    """Добавить заголовок с нужным стилем."""
    sizes = {1: 16, 2: 14, 3: 13}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(sizes.get(level, 12))
    run.font.name = 'Times New Roman'
    return p


def add_bullet_paragraph(doc, text, size=12, indent=True):
    """Добавить маркированный пункт."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = 1.5
    if indent:
        p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run('• ' + text)
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p


def add_numbered_paragraph(doc, text, number, size=12):
    """Добавить нумерованный пункт."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(1.25)
    run = p.add_run(f'{number}. {text}')
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return p


def set_margins(doc, top=2, bottom=2, left=3, right=1.5):
    """Установить поля документа (в см)."""
    for section in doc.sections:
        section.top_margin = Cm(top)
        section.bottom_margin = Cm(bottom)
        section.left_margin = Cm(left)
        section.right_margin = Cm(right)


def create_note():
    doc = Document()

    # Настройка полей
    set_margins(doc)

    # Настройка стиля по умолчанию
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5

    # ==================== ТИТУЛЬНЫЙ ЛИСТ ====================
    for _ in range(6):
        doc.add_paragraph()

    add_formatted_paragraph(doc, 'ПОЯСНИТЕЛЬНАЯ ЗАПИСКА', bold=True, size=16,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    add_formatted_paragraph(doc, 'к дипломному проекту на тему:', size=14,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=24)
    add_formatted_paragraph(doc, '«Разработка интернет-магазина цифровых товаров DigitalStore»',
                            bold=True, size=14, alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=36)

    for _ in range(8):
        doc.add_paragraph()

    add_formatted_paragraph(doc, 'Выполнил: студент(ка) _______________', size=12,
                            alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)
    add_formatted_paragraph(doc, 'Проверил: преподаватель _______________', size=12,
                            alignment=WD_ALIGN_PARAGRAPH.RIGHT, space_after=6)

    for _ in range(4):
        doc.add_paragraph()

    year = datetime.datetime.now().year
    add_formatted_paragraph(doc, f'Москва, {year}', size=12,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_page_break()

    # ==================== СОДЕРЖАНИЕ ====================
    add_formatted_paragraph(doc, 'СОДЕРЖАНИЕ', bold=True, size=14,
                            alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    toc_items = [
        'Введение',
        '1. Анализ предметной области',
        '   1.1. Обзор существующих решений',
        '   1.2. Постановка задачи',
        '   1.3. Требования к системе',
        '2. Проектирование системы',
        '   2.1. Архитектура приложения',
        '   2.2. Проектирование базы данных',
        '   2.3. Проектирование пользовательского интерфейса',
        '3. Реализация системы',
        '   3.1. Среда разработки и технологический стек',
        '   3.2. Реализация модуля каталога товаров',
        '   3.3. Реализация модуля корзины',
        '   3.4. Реализация модуля заказов и оплаты',
        '   3.5. Реализация модуля скачивания цифровых товаров',
        '   3.6. Реализация модуля управления пользователями',
        '   3.7. Административная панель',
        '4. Тестирование и развёртывание',
        '   4.1. Виды тестирования',
        '   4.2. Развёртывание приложения',
        'Заключение',
        'Список использованных источников',
        'Приложение А. Структура базы данных (ER-диаграмма)',
        'Приложение Б. Основные фрагменты исходного кода',
    ]

    for item in toc_items:
        add_formatted_paragraph(doc, item, size=12, space_after=3)

    doc.add_page_break()

    # ==================== ВВЕДЕНИЕ ====================
    add_heading_styled(doc, 'ВВЕДЕНИЕ', level=1)

    add_formatted_paragraph(doc,
        'Современная электронная коммерция охватывает всё больше сфер человеческой деятельности. '
        'Одним из наиболее динамично развивающихся сегментов является рынок цифровых товаров — '
        'программных продуктов, электронных книг, графических шаблонов, аудио- и видеоматериалов, '
        'обучающих курсов и иных продуктов, распространяемых в электронном виде.', size=12)

    add_formatted_paragraph(doc,
        'Разработка специализированного интернет-магазина цифровых товаров представляет собой '
        'актуальную задачу, поскольку такой магазин должен обеспечивать не только каталогизацию '
        'и продажу продукции, но и механизм безопасной доставки цифрового контента покупателю, '
        'контроль сроков доступа к скачиванию, а также интеграцию с платёжными системами.', size=12)

    add_formatted_paragraph(doc,
        'Целью данного дипломного проекта является проектирование и разработка полнофункционального '
        'веб-приложения интернет-магазина цифровых товаров «DigitalStore» на основе фреймворка Django.', size=12)

    add_formatted_paragraph(doc,
        'Для достижения поставленной цели необходимо решить следующие задачи:', size=12)

    add_bullet_paragraph(doc, 'проанализировать предметную область и существующие решения;')
    add_bullet_paragraph(doc, 'сформулировать функциональные и нефункциональные требования к системе;')
    add_bullet_paragraph(doc, 'спроектировать архитектуру приложения и структуру базы данных;')
    add_bullet_paragraph(doc, 'реализовать модули каталога, корзины, заказов, оплаты, скачивания и управления пользователями;')
    add_bullet_paragraph(doc, 'разработать административную панель для управления контентом;')
    add_bullet_paragraph(doc, 'провести тестирование и подготовить приложение к развёртыванию.')

    add_formatted_paragraph(doc,
        'В результате выполнения работы было создано веб-приложение, позволяющее пользователям '
        'просматривать каталог цифровых товаров, формировать корзину покупок, оформлять заказы, '
        'оплачивать их и получать безопасные ссылки для скачивания приобретённых продуктов.', size=12)

    add_formatted_paragraph(doc,
        'Документ состоит из введения, четырёх глав, заключения, списка использованных источников '
        'и приложений, содержащих структуру базы данных и ключевые фрагменты исходного кода.', size=12)

    doc.add_page_break()

    # ==================== ГЛАВА 1 ====================
    add_heading_styled(doc, '1. АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ', level=1)

    # 1.1
    add_heading_styled(doc, '1.1. Обзор существующих решений', level=2)

    add_formatted_paragraph(doc,
        'На современном рынке электронной коммерции представлено множество платформ для продажи '
        'цифровых товаров. Рассмотрим наиболее значимые из них.', size=12)

    add_heading_styled(doc, 'Gumroad', level=3)
    add_formatted_paragraph(doc,
        'Gumroad — одна из наиболее популярных платформ для продажи цифровых товаров. '
        'Позволяет авторам продавать электронные книги, курсы, программное обеспечение, '
        'музыку и другие цифровые продукты. Платформа берёт комиссию с каждой продажи. '
        'Преимущества: простота настройки, встроенная аналитика, поддержка подписок. '
        'Недостатки: ограниченная кастомизация, комиссия с продаж, зависимость от платформы.', size=12)

    add_heading_styled(doc, 'Sellfy', level=3)
    add_formatted_paragraph(doc,
        'Sellfy — платформа электронной коммерции, ориентированная на продажу цифровых, '
        'физических товаров и подписок. Предлагает встроенные маркетинговые инструменты, '
        'интеграцию с социальными сетями. Недостатки: ограниченный бесплатный тариф, '
        'отсутствие полноценной локализации для русскоязычного рынка.', size=12)

    add_heading_styled(doc, 'WooCommerce + плагины', level=3)
    add_formatted_paragraph(doc,
        'WooCommerce — плагин для WordPress, позволяющий создать интернет-магазин. '
        'С помощью дополнительных плагинов (например, WooCommerce Downloadable Products) '
        'можно организовать продажу цифровых товаров. Преимущества: гибкость, большое '
        'количество плагинов, открытый исходный код. Недостатки: зависимость от WordPress, '
        'необходимость самостоятельного хостинга и обслуживания, потенциальные проблемы '
        'с безопасностью.', size=12)

    add_heading_styled(doc, 'OpenCart', level=3)
    add_formatted_paragraph(doc,
        'OpenCart — свободная платформа электронной коммерции на PHP. Поддерживает продажу '
        'цифровых товаров через дополнительные модули. Преимущества: открытый код, '
        'большое сообщество, множество тем и модулей. Недостатки: устаревшая архитектура, '
        'ограниченная поддержка современных практик разработки.', size=12)

    add_formatted_paragraph(doc,
        'Анализ существующих решений показал, что каждое из них имеет определённые ограничения: '
        'либо зависимость от сторонней платформы (Gumroad, Sellfy), либо избыточная сложность '
        'и зависимость от CMS (WooCommerce), либо устаревшая архитектура (OpenCart). '
        'Это обосновывает необходимость разработки специализированного решения на базе '
        'современного фреймворка Django, обеспечивающего полный контроль над функциональностью, '
        'безопасностью и масштабируемостью приложения.', size=12)

    # 1.2
    add_heading_styled(doc, '1.2. Постановка задачи', level=2)

    add_formatted_paragraph(doc,
        'На основании проведённого анализа предметной области сформулирована следующая задача: '
        'разработать веб-приложение интернет-магазина цифровых товаров «DigitalStore», '
        'обеспечивающее полный цикл работы с цифровым продуктом — от каталогизации до '
        'доставки покупателю.', size=12)

    add_formatted_paragraph(doc, 'Основные функциональные требования:', size=12)

    add_numbered_paragraph(doc,
        'Каталог товаров с категориями, фильтрацией, поиском и пагинацией', 1)
    add_numbered_paragraph(doc,
        'Карточка товара с описанием, изображением предпросмотра и связанными товарами', 2)
    add_numbered_paragraph(doc,
        'Корзина покупок с возможностью добавления, удаления и изменения количества товаров', 3)
    add_numbered_paragraph(doc,
        'Оформление заказа с расчётом итоговой стоимости', 4)
    add_numbered_paragraph(doc,
        'Интеграция с платёжной системой (заглушка с поддержкой webhook)', 5)
    add_numbered_paragraph(doc,
        'Система скачивания цифровых товаров через уникальные одноразовые ссылки с ограниченным сроком действия', 6)
    add_numbered_paragraph(doc,
        'Регистрация, авторизация и управление профилем пользователя', 7)
    add_numbered_paragraph(doc,
        'История заказов и список доступных загрузок в личном кабинете', 8)
    add_numbered_paragraph(doc,
        'Административная панель для управления товарами, категориями, заказами и загрузками', 9)

    add_formatted_paragraph(doc, 'Нефункциональные требования:', size=12)

    add_bullet_paragraph(doc, 'Адаптивный пользовательский интерфейс на основе Bootstrap 5;')
    add_bullet_paragraph(doc, 'Защита от CSRF-атак, SQL-инъекций и XSS;')
    add_bullet_paragraph(doc, 'Локализация интерфейса на русском языке;')
    add_bullet_paragraph(doc, 'Использование реляционной СУБД SQLite (с возможностью миграции на PostgreSQL);')
    add_bullet_paragraph(doc, 'Модульная архитектура с разделением ответственности между приложениями Django.')

    # 1.3
    add_heading_styled(doc, '1.3. Требования к системе', level=2)

    add_formatted_paragraph(doc,
        'На основании постановки задачи составлены детализированные требования к системе '
        'в виде диаграмм вариантов использования (Use Case) и описания ключевых сценариев.', size=12)

    add_formatted_paragraph(doc, 'Основные акторы системы:', size=12)

    add_bullet_paragraph(doc, '<b>Гость</b> — неавторизованный посетитель, который может просматривать каталог, '
                   'искать товары и просматривать карточки продуктов.')
    add_bullet_paragraph(doc, '<b>Покупатель</b> — зарегистрированный и авторизованный пользователь, который может '
                   'добавлять товары в корзину, оформлять заказы, оплачивать их и скачивать приобретённые продукты.')
    add_bullet_paragraph(doc, '<b>Администратор</b> — пользователь с правами управления контентом магазина через '
                   'встроенную административную панель Django.')

    add_formatted_paragraph(doc, 'Ключевые сценарии использования:', size=12)

    # Таблица сценариев
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['№', 'Сценарий', 'Описание']
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)

    scenarios = [
        ('1', 'Просмотр каталога', 'Пользователь просматривает список товаров с возможностью фильтрации по категории, цене и поисковому запросу'),
        ('2', 'Просмотр товара', 'Пользователь открывает карточку товара с подробным описанием, изображением и рекомендациями'),
        ('3', 'Добавление в корзину', 'Пользователь добавляет товар в корзину (с указанием количества) через AJAX-запрос'),
        ('4', 'Оформление заказа', 'Авторизованный пользователь переходит к оформлению заказа из корзины'),
        ('5', 'Оплата заказа', 'Пользователь оплачивает заказ через интегрированную платёжную систему'),
        ('6', 'Скачивание товара', 'После оплаты пользователь получает уникальную одноразовую ссылку для скачивания'),
        ('7', 'Управление профилем', 'Пользователь редактирует свои данные (email, телефон) в личном кабинете'),
        ('8', 'Администрирование', 'Администратор управляет товарами, категориями, заказами через панель управления'),
    ]

    for num, name, desc in scenarios:
        row = table.add_row()
        row.cells[0].text = num
        row.cells[1].text = name
        row.cells[2].text = desc
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    add_formatted_paragraph(doc, '', size=12, space_after=6)

    doc.add_page_break()

    # ==================== ГЛАВА 2 ====================
    add_heading_styled(doc, '2. ПРОЕКТИРОВАНИЕ СИСТЕМЫ', level=1)

    # 2.1
    add_heading_styled(doc, '2.1. Архитектура приложения', level=2)

    add_formatted_paragraph(doc,
        'Приложение DigitalStore разработано на основе фреймворка Django версии 4.2+ и следует '
        'архитектурному паттерну MVT (Model-View-Template), который является вариацией классического '
        'паттерна MVC (Model-View-Controller).', size=12)

    add_formatted_paragraph(doc, 'Компоненты архитектуры MVT:', size=12)

    add_bullet_paragraph(doc, '<b>Model (Модель)</b> — определяет структуру данных и взаимодействие с базой данных. '
                   'Реализована через ORM Django. Модели: Category, Product, Order, OrderItem, Download, Profile.')
    add_bullet_paragraph(doc, '<b>View (Представление)</b> — содержит бизнес-логику приложения, обрабатывает HTTP-запросы '
                   'и формирует ответы. Используются как классовые представления (ListView, DetailView), '
                   'так и функциональные.')
    add_bullet_paragraph(doc, '<b>Template (Шаблон)</b> — отвечает за представление данных пользователю. Шаблоны написаны '
                   'на HTML с использованием шаблонизатора Django (Django Template Language) и фреймворка Bootstrap 5.')

    add_formatted_paragraph(doc,
        'Приложение разделено на пять модулей (приложений Django), каждый из которых отвечает '
        'за определённую функциональную область:', size=12)

    # Таблица модулей
    table2 = doc.add_table(rows=1, cols=3)
    table2.style = 'Table Grid'
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers2 = ['Приложение', 'Назначение', 'Ключевые модели']
    for i, header in enumerate(headers2):
        cell = table2.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)

    modules = [
        ('products', 'Каталог товаров, категории, поиск', 'Category, Product'),
        ('cart', 'Корзина покупок (сессионная)', 'Cart (класс, не модель)'),
        ('orders', 'Заказы, позиции заказа, загрузки', 'Order, OrderItem, Download'),
        ('users', 'Регистрация, авторизация, профиль', 'Profile'),
        ('payments', 'Создание и обработка платежей', '— (использует Order)'),
    ]

    for app, purpose, models in modules:
        row = table2.add_row()
        row.cells[0].text = app
        row.cells[1].text = purpose
        row.cells[2].text = models
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    add_formatted_paragraph(doc, '', size=12, space_after=6)

    add_formatted_paragraph(doc,
        'Взаимодействие модулей организовано следующим образом: модуль products предоставляет '
        'данные о товарах для всех остальных модулей; cart использует данные Product для '
        'формирования корзины; orders создаёт заказы на основе содержимого корзины; payments '
        'обрабатывает оплату заказов; users управляет аутентификацией и профилями. Модуль cart '
        'реализован как сессионная корзина — данные хранятся в сессии пользователя и не требуют '
        'авторизации для добавления товаров.', size=12)

    # 2.2
    add_heading_styled(doc, '2.2. Проектирование базы данных', level=2)

    add_formatted_paragraph(doc,
        'Для хранения данных приложения используется реляционная СУБД SQLite, входящая в состав '
        'стандартной поставки Django. При необходимости система может быть легко мигрирована на '
        'PostgreSQL путём изменения конфигурации в файле settings.py.', size=12)

    add_formatted_paragraph(doc,
        'Структура базы данных включает следующие основные сущности:', size=12)

    add_heading_styled(doc, 'Модель Category', level=3)
    add_formatted_paragraph(doc,
        'Хранит информацию о категориях товаров. Поля: name (название), slug (уникальный '
        'URL-идентификатор), description (описание), created_at (дата создания). Связь «один-ко-многим» '
        'с моделью Product.', size=12)

    add_heading_styled(doc, 'Модель Product', level=3)
    add_formatted_paragraph(doc,
        'Хранит информацию о товарах. Поля: title (название), slug (уникальный URL-идентификатор, '
        'генерируется автоматически с транслитерацией кириллицы), description (описание), price (цена), '
        'category (внешний ключ на Category), preview_image (изображение предпросмотра), digital_file '
        '(цифровой файл для скачивания), created_at, updated_at, is_active (флаг активности).', size=12)

    add_heading_styled(doc, 'Модель Order', level=3)
    add_formatted_paragraph(doc,
        'Хранит информацию о заказах. Поля: user (внешний ключ на User), total_price (общая стоимость), '
        'status (статус: pending/paid/cancelled), payment_id (идентификатор платежа), created_at, updated_at. '
        'Связь «один-ко-многим» с OrderItem и Download.', size=12)

    add_heading_styled(doc, 'Модель OrderItem', level=3)
    add_formatted_paragraph(doc,
        'Хранит позиции заказа. Поля: order (внешний ключ на Order), product (внешний ключ на Product), '
        'price (цена на момент заказа), quantity (количество). Вычисляемое свойство total = price × quantity.', size=12)

    add_heading_styled(doc, 'Модель Download', level=3)
    add_formatted_paragraph(doc,
        'Хранит информацию о ссылках для скачивания. Поля: user, product, order (внешние ключи), '
        'download_token (уникальный UUID-токен), created_at, expires_at (срок действия — 30 дней), '
        'is_used (флаг использования — ссылка одноразовая).', size=12)

    add_heading_styled(doc, 'Модель Profile', level=3)
    add_formatted_paragraph(doc,
        'Расширяет стандартную модель пользователя Django. Поля: user (OneToOne на auth.User), '
        'phone (телефон), created_at.', size=12)

    add_formatted_paragraph(doc,
        'Связи между моделями организованы следующим образом: Category → Product (один-ко-многим), '
        'User → Order (один-ко-многим), Order → OrderItem (один-ко-многим), Order → Download '
        '(один-ко-многим), User → Profile (один-к-одному). Все внешние ключи настроены с каскадным '
        'удалением (on_delete=CASCADE).', size=12)

    # 2.3
    add_heading_styled(doc, '2.3. Проектирование пользовательского интерфейса', level=2)

    add_formatted_paragraph(doc,
        'Пользовательский интерфейс разработан с использованием HTML5, CSS3 (фреймворк Bootstrap 5) '
        'и JavaScript. Шаблоны построены на основе наследования: базовый шаблон base.html определяет '
        'общую структуру страниц (навигация, футер, подключение стилей), а дочерние шаблоны '
        'переопределяют блок content.', size=12)

    add_formatted_paragraph(doc, 'Основные страницы приложения:', size=12)

    add_bullet_paragraph(doc, '<b>Главная страница</b> — отображает список категорий товаров.')
    add_bullet_paragraph(doc, '<b>Каталог</b> — пагинированный список товаров с фильтрами (категория, цена, поиск).')
    add_bullet_paragraph(doc, '<b>Карточка товара</b> — подробное описание, изображение, кнопка «В корзину», похожие товары.')
    add_bullet_paragraph(doc, '<b>Корзина</b> — список добавленных товаров с возможностью изменения количества и удаления.')
    add_bullet_paragraph(doc, '<b>Оформление заказа</b> — подтверждение заказа и переход к оплате.')
    add_bullet_paragraph(doc, '<b>Оплата</b> — страница оплаты с перенаправлением на успех/неудачу.')
    add_bullet_paragraph(doc, '<b>История заказов</b> — список оплаченных заказов пользователя.')
    add_bullet_paragraph(doc, '<b>Профиль</b> — данные пользователя, доступные загрузки, история заказов.')
    add_bullet_paragraph(doc, '<b>Регистрация / Авторизация</b> — формы входа и регистрации с валидацией.')

    add_formatted_paragraph(doc,
        'Для обеспечения интерактивности без перезагрузки страницы используются AJAX-запросы '
        'при добавлении товаров в корзину и их удалении. Библиотека django-crispy-forms с пакетом '
        'Bootstrap 4 используется для стилизации форм.', size=12)

    doc.add_page_break()

    # ==================== ГЛАВА 3 ====================
    add_heading_styled(doc, '3. РЕАЛИЗАЦИЯ СИСТЕМЫ', level=1)

    # 3.1
    add_heading_styled(doc, '3.1. Среда разработки и технологический стек', level=2)

    add_formatted_paragraph(doc,
        'Разработка приложения велась в следующей среде с использованием указанных технологий:', size=12)

    # Таблица технологий
    table3 = doc.add_table(rows=1, cols=2)
    table3.style = 'Table Grid'
    table3.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers3 = ['Технология', 'Назначение']
    for i, header in enumerate(headers3):
        cell = table3.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(11)

    techs = [
        ('Python 3.9+', 'Язык программирования'),
        ('Django 4.2+', 'Веб-фреймворк'),
        ('SQLite', 'Система управления базами данных'),
        ('Bootstrap 5', 'CSS-фреймворк для интерфейса'),
        ('django-crispy-forms + crispy-bootstrap4', 'Стилизация форм'),
        ('Pillow', 'Обработка изображений (ImageField)'),
        ('python-dotenv', 'Управление переменными окружения'),
        ('JavaScript (AJAX)', 'Асинхронные запросы к серверу'),
        ('HTML5 / Django Template Language', 'Шаблонизация страниц'),
        ('Git', 'Система контроля версий'),
    ]

    for tech, purpose in techs:
        row = table3.add_row()
        row.cells[0].text = tech
        row.cells[1].text = purpose
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    add_formatted_paragraph(doc, '', size=12, space_after=6)

    add_formatted_paragraph(doc,
        'Управление зависимостями осуществляется через файл requirements.txt. Конфигурация приложения '
        'поддерживает загрузку переменных окружения из файла .env с помощью библиотеки python-dotenv.', size=12)

    # 3.2
    add_heading_styled(doc, '3.2. Реализация модуля каталога товаров', level=2)

    add_formatted_paragraph(doc,
        'Модуль каталога реализован в приложении products и включает следующие компоненты:', size=12)

    add_formatted_paragraph(doc, 'Модели данных:', size=12)
    add_bullet_paragraph(doc, '<b>Category</b> — модель категории с полями name, slug, description, created_at. '
                   'Поддержка сортировки по названию.')
    add_bullet_paragraph(doc, '<b>Product</b> — модель товара с полями title, slug, description, price, category, '
                   'preview_image, digital_file, created_at, updated_at, is_active. Автоматическая генерация '
                   'slug с транслитерацией кириллицы при сохранении.')

    add_formatted_paragraph(doc, 'Представления:', size=12)
    add_bullet_paragraph(doc, '<b>home</b> — функциональное представление главной страницы, отображает список категорий.')
    add_bullet_paragraph(doc, '<b>CatalogView</b> — классовое представление ListView с пагинацией (12 товаров на страницу), '
                   'фильтрацией по категории, диапазону цен и поисковому запросу.')
    add_bullet_paragraph(doc, '<b>ProductDetailView</b> — классовое представление DetailView, отображает карточку товара '
                   'с четырьмя связанными товарами из той же категории.')
    add_bullet_paragraph(doc, '<b>search_products</b> — функциональное представление для поиска по названию и описанию.')

    add_formatted_paragraph(doc,
        'Генерация slug реализована функцией generate_unique_slug(), которая выполняет транслитерацию '
        'кириллических символов в латиницу, применяет стандартную функцию Django slugify() и обеспечивает '
        'уникальность путём добавления числового суффикса при необходимости. Если транслитерация даёт '
        'пустую строку, используется UUID.', size=12)

    # 3.3
    add_heading_styled(doc, '3.3. Реализация модуля корзины', level=2)

    add_formatted_paragraph(doc,
        'Модуль корзины реализован в приложении cart. Корзина является сессионной — данные хранятся '
        'в сессии пользователя (request.session) и не требуют авторизации для добавления товаров.', size=12)

    add_formatted_paragraph(doc,
        'Основной класс Cart содержит следующие методы:', size=12)

    add_bullet_paragraph(doc, '<b>add(product, quantity)</b> — добавляет товар в корзину или увеличивает его количество.')
    add_bullet_paragraph(doc, '<b>remove(product)</b> — удаляет товар из корзины.')
    add_bullet_paragraph(doc, '<b>update(product, quantity)</b> — обновляет количество товара (удаляет при quantity ≤ 0).')
    add_bullet_paragraph(doc, '<b>__iter__()</b> — итератор, обогащающий элементы корзины объектами Product, ценами и итогами.')
    add_bullet_paragraph(doc, '<b>__len__()</b> — возвращает общее количество единиц товаров в корзине.')
    add_bullet_paragraph(doc, '<b>get_total_price()</b> — возвращает общую стоимость корзины.')
    add_bullet_paragraph(doc, '<b>clear()</b> — очищает корзину.')

    add_formatted_paragraph(doc,
        'Контекстный процессор cart_processor обеспечивает доступ к объекту корзины и количеству товаров '
        'во всех шаблонах без явной передачи в каждом представлении. URL-маршруты модуля поддерживают '
        'операции добавления, удаления и обновления товаров по их идентификаторам.', size=12)

    # 3.4
    add_heading_styled(doc, '3.4. Реализация модуля заказов и оплаты', level=2)

    add_formatted_paragraph(doc,
        'Модуль заказов реализован в приложении orders. Процесс оформления заказа включает следующие этапы:', size=12)

    add_numbered_paragraph(doc,
        'Пользователь переходит к оформлению из корзины (представление checkout)', 1)
    add_numbered_paragraph(doc,
        'Создаётся объект Order со статусом «pending» (ожидает оплаты)', 2)
    add_numbered_paragraph(doc,
        'Для каждого элемента корзины создаётся OrderItem с фиксацией цены на момент заказа', 3)
    add_numbered_paragraph(doc,
        'Корзина очищается, пользователь перенаправляется на страницу оплаты', 4)
    add_numbered_paragraph(doc,
        'После успешной оплаты статус заказа меняется на «paid», создаются записи Download', 5)

    add_formatted_paragraph(doc,
        'Модуль оплаты (payments) реализован в виде заглушки, имитирующей работу платёжной системы. '
        'Представление payment_create отображает страницу оплаты, payment_success обрабатывает успешную '
        'оплату и создаёт ссылки для скачивания, payment_fail — неудачную оплату. Функция payment_webhook '
        'принимает POST-запросы от платёжной системы в формате, совместимом с YooKassa, и обновляет '
        'статус заказа при получении уведомления об успешном платеже.', size=12)

    # 3.5
    add_heading_styled(doc, '3.5. Реализация модуля скачивания цифровых товаров', level=2)

    add_formatted_paragraph(doc,
        'Система скачивания цифровых товаров обеспечивает безопасную доставку контента покупателю. '
        'Механизм работает следующим образом:', size=12)

    add_numbered_paragraph(doc,
        'После успешной оплаты заказа для каждой позиции заказа создаётся запись Download', 1)
    add_numbered_paragraph(doc,
        'Каждая запись содержит уникальный UUID-токен (download_token), который используется в URL', 2)
    add_numbered_paragraph(doc,
        'Устанавливается срок действия ссылки (expires_at) — 30 дней с момента создания', 3)
    add_numbered_paragraph(doc,
        'При скачивании проверяется: не использована ли ссылка (is_used), не истёк ли срок (expires_at)', 4)
    add_numbered_paragraph(doc,
        'После успешного скачивания ссылка помечается как использованная (is_used = True)', 5)
    add_numbered_paragraph(doc,
        'Файл отдаётся через Django FileResponse как вложение (as_attachment=True)', 6)

    add_formatted_paragraph(doc,
        'Данный подход обеспечивает контроль доступа к цифровому контенту: ссылка является одноразовой, '
        'имеет ограниченный срок действия и привязана к конкретному пользователю. Попытка повторного '
        'скачивания или использования просроченной ссылки блокируется с выводом соответствующего '
        'сообщения об ошибке.', size=12)

    # 3.6
    add_heading_styled(doc, '3.6. Реализация модуля управления пользователями', level=2)

    add_formatted_paragraph(doc,
        'Модуль управления пользователями реализован в приложении users и включает следующие функции:', size=12)

    add_bullet_paragraph(doc, '<b>Регистрация</b> — форма RegisterForm на основе UserCreationForm с обязательным полем email. '
                   'После регистрации автоматически создаётся профиль (Profile) и выполняется вход в систему.')
    add_bullet_paragraph(doc, '<b>Авторизация</b> — стандартная форма AuthenticationForm Django с обработкой ошибок.')
    add_bullet_paragraph(doc, '<b>Выход</b> — стандартная функция logout Django с перенаправлением на главную страницу.')
    add_bullet_paragraph(doc, '<b>Профиль</b> — отображает список доступных загрузок и историю заказов.')
    add_bullet_paragraph(doc, '<b>Редактирование профиля</b> — форма ProfileForm для изменения номера телефона.')

    add_formatted_paragraph(doc,
        'Модель Profile расширяет стандартную модель пользователя Django (auth.User) через отношение '
        'OneToOneField. Это позволяет добавить дополнительные поля (телефон) без переопределения '
        'встроенной модели пользователя, что является рекомендуемой практикой Django.', size=12)

    # 3.7
    add_heading_styled(doc, '3.7. Административная панель', level=2)

    add_formatted_paragraph(doc,
        'Для управления контентом магазина используется встроенная административная панель Django, '
        'настроенная для каждой модели:', size=12)

    add_bullet_paragraph(doc, '<b>CategoryAdmin</b> — автозаполнение slug из названия, поиск по имени, фильтрация по дате.')
    add_bullet_paragraph(doc, '<b>ProductAdmin</b> — автозаполнение slug, редактирование флага is_active из списка, '
                   'фильтрация по категории и активности, иерархия по дате.')
    add_bullet_paragraph(doc, '<b>OrderAdmin</b> — встроенные inline-формы для OrderItem и Download, фильтрация по статусу, '
                   'поиск по имени пользователя и идентификатору платежа.')
    add_bullet_paragraph(doc, '<b>DownloadAdmin</b> — отображение информации о загрузках, фильтрация по статусу использования.')

    add_formatted_paragraph(doc,
        'Административная панель доступна по адресу /admin/ и предоставляет полный контроль над '
        'всеми данными магазина без необходимости разработки дополнительного интерфейса.', size=12)

    doc.add_page_break()

    # ==================== ГЛАВА 4 ====================
    add_heading_styled(doc, '4. ТЕСТИРОВАНИЕ И РАЗВЁРТЫВАНИЕ', level=1)

    # 4.1
    add_heading_styled(doc, '4.1. Виды тестирования', level=2)

    add_formatted_paragraph(doc,
        'В процессе разработки приложения были проведены следующие виды тестирования:', size=12)

    add_heading_styled(doc, 'Функциональное тестирование', level=3)
    add_formatted_paragraph(doc,
        'Проверка корректности работы всех функций приложения: регистрация и авторизация пользователей, '
        'навигация по каталогу, добавление товаров в корзину, оформление заказа, оплата, скачивание '
        'файлов. Все основные сценарии использования протестированы вручную.', size=12)

    add_heading_styled(doc, 'Тестирование пользовательского интерфейса', level=3)
    add_formatted_paragraph(doc,
        'Проверка корректности отображения страниц в различных браузерах и на различных разрешениях '
        'экрана. Адаптивность интерфейса обеспечена фреймворком Bootstrap 5.', size=12)

    add_heading_styled(doc, 'Тестирование безопасности', level=3)
    add_formatted_paragraph(doc,
        'Проверена защита от основных типов атак: CSRF-защита включена по умолчанию через промежуточное '
        'ПО CsrfViewMiddleware; SQL-инъекции предотвращены использованием ORM Django; XSS-атаки — '
        'автоматическим экранированием в шаблонах Django. Доступ к скачиванию файлов защищён '
        'аутентификацией и проверкой прав пользователя.', size=12)

    add_heading_styled(doc, 'Тестирование граничных случаев', level=3)
    add_formatted_paragraph(doc,
        'Проверены граничные ситуации: попытка скачивания по использованной ссылке, попытка скачивания '
        'по просроченной ссылке, оформление заказа с пустой корзиной, доступ к страницам, требующим '
        'авторизации, без входа в систему. Во всех случаях приложение корректно обрабатывает ситуации '
        'и выводит информативные сообщения об ошибках.', size=12)

    # 4.2
    add_heading_styled(doc, '4.2. Развёртывание приложения', level=2)

    add_formatted_paragraph(doc,
        'Для развёртывания приложения необходимо выполнить следующие шаги:', size=12)

    add_numbered_paragraph(doc, 'Клонировать репозиторий проекта', 1)
    add_numbered_paragraph(doc, 'Создать и активировать виртуальное окружение Python', 2)
    add_numbered_paragraph(doc, 'Установить зависимости: pip install -r requirements.txt', 3)
    add_numbered_paragraph(doc, 'Создать файл .env с переменными SECRET_KEY, DEBUG, ALLOWED_HOSTS', 4)
    add_numbered_paragraph(doc, 'Выполнить миграции: python manage.py migrate', 5)
    add_numbered_paragraph(doc, 'Создать суперпользователя: python manage.py createsuperuser', 6)
    add_numbered_paragraph(doc, 'Запустить сервер: python manage.py runserver', 7)

    add_formatted_paragraph(doc,
        'Для production-развёртывания рекомендуется использовать WSGI-сервер Gunicorn совместно с '
        'веб-сервером Nginx, настроить переменную DEBUG=False, задать безопасный SECRET_KEY, '
        'сконфигурировать ALLOWED_HOSTS и перейти на СУБД PostgreSQL.', size=12)

    doc.add_page_break()

    # ==================== ЗАКЛЮЧЕНИЕ ====================
    add_heading_styled(doc, 'ЗАКЛЮЧЕНИЕ', level=1)

    add_formatted_paragraph(doc,
        'В рамках дипломного проекта было разработано полнофункциональное веб-приложение '
        'интернет-магазина цифровых товаров «DigitalStore» на основе фреймворка Django.', size=12)

    add_formatted_paragraph(doc,
        'В ходе выполнения работы были решены все поставленные задачи:', size=12)

    add_bullet_paragraph(doc, 'Проведён анализ предметной области и существующих решений для продажи цифровых товаров.')
    add_bullet_paragraph(doc, 'Сформулированы функциональные и нефункциональные требования к системе.')
    add_bullet_paragraph(doc, 'Спроектирована архитектура приложения на основе паттерна MVT и структура базы данных.')
    add_bullet_paragraph(doc, 'Реализован модуль каталога товаров с категориями, фильтрацией, поиском и пагинацией.')
    add_bullet_paragraph(doc, 'Реализован модуль сессионной корзины покупок с AJAX-взаимодействием.')
    add_bullet_paragraph(doc, 'Реализован модуль заказов с поддержкой статусов и историей покупок.')
    add_bullet_paragraph(doc, 'Реализован модуль оплаты с имитацией работы платёжной системы и webhook-обработчиком.')
    add_bullet_paragraph(doc, 'Реализована безопасная система скачивания цифровых товаров через одноразовые '
                   'UUID-ссылки с ограниченным сроком действия.')
    add_bullet_paragraph(doc, 'Реализован модуль управления пользователями с регистрацией, авторизацией и профилем.')
    add_bullet_paragraph(doc, 'Настроена административная панель для управления всеми данными магазина.')

    add_formatted_paragraph(doc,
        'Результатом работы является работоспособное веб-приложение, которое может быть использовано '
        'для продажи цифровых товаров. Архитектура приложения позволяет легко расширять функциональность: '
        'добавлять новые способы оплаты, интегрировать полноценную платёжную систему, реализовывать '
        'систему скидок и промокодов, добавлять отзывы к товарам.', size=12)

    doc.add_page_break()

    # ==================== СПИСОК ИСТОЧНИКОВ ====================
    add_heading_styled(doc, 'СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ', level=1)

    sources = [
        'Документация Django 4.2 [Электронный ресурс]. — URL: https://docs.djangoproject.com/en/4.2/ (дата обращения: 01.03.2026).',
        'Документация Bootstrap 5 [Электронный ресурс]. — URL: https://getbootstrap.com/docs/5.0/ (дата обращения: 01.03.2026).',
        'Forbes A., Bussell D. The Definitive Guide to Django: Web Development Done Right. — Apress, 2007. — 496 p.',
        'Holovaty A., Kaplan-Moss J. The Django Book. — Apress, 2020. — 650 p.',
        'Документация django-crispy-forms [Электронный ресурс]. — URL: https://django-crispy-forms.readthedocs.io/ (дата обращения: 01.03.2026).',
        'OWASP Top Ten Web Application Security Risks [Электронный ресурс]. — URL: https://owasp.org/www-project-top-ten/ (дата обращения: 01.03.2026).',
        'Документация YooKassa для разработчиков [Электронный ресурс]. — URL: https://yookassa.ru/developers (дата обращения: 01.03.2026).',
        'Lutz M. Learning Python. — 5th Edition. — O\'Reilly Media, 2013. — 1648 p.',
    ]

    for source in sources:
        add_formatted_paragraph(doc, source, size=12, space_after=4)

    doc.add_page_break()

    # ==================== ПРИЛОЖЕНИЕ А ====================
    add_heading_styled(doc, 'ПРИЛОЖЕНИЕ А. Структура базы данных (ER-диаграмма)', level=1)

    add_formatted_paragraph(doc,
        'Ниже представлено текстовое описание ER-диаграммы базы данных приложения DigitalStore. '
        'Система использует реляционную модель данных со следующими сущностями и связями:', size=12)

    add_formatted_paragraph(doc, 'Сущности и их атрибуты:', size=12)

    # Таблица ER
    table_er = doc.add_table(rows=1, cols=3)
    table_er.style = 'Table Grid'
    table_er.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers_er = ['Сущность', 'Атрибуты', 'Тип']
    for i, header in enumerate(headers_er):
        cell = table_er.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, 'D9E2F3')
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)

    er_data = [
        ('Category', 'id, name, slug, description, created_at', 'PK, Char, Slug, Text, DateTime'),
        ('Product', 'id, title, slug, description, price, category_id, preview_image, digital_file, created_at, updated_at, is_active', 'PK, Char, Slug, Text, Decimal, FK, Image, File, DateTime, DateTime, Boolean'),
        ('auth.User', 'id, username, email, password, first_name, last_name, is_staff, is_active, date_joined', 'PK, Char, Email, Char, Char, Char, Boolean, Boolean, DateTime'),
        ('Profile', 'id, user_id, phone, created_at', 'PK, OneToOne(FK), Char, DateTime'),
        ('Order', 'id, user_id, total_price, status, payment_id, created_at, updated_at', 'PK, FK, Decimal, Char, Char, DateTime, DateTime'),
        ('OrderItem', 'id, order_id, product_id, price, quantity', 'PK, FK, FK, Decimal, PositiveInt'),
        ('Download', 'id, user_id, product_id, order_id, download_token, created_at, expires_at, is_used', 'PK, FK, FK, FK, UUID, DateTime, DateTime, Boolean'),
    ]

    for entity, attrs, types in er_data:
        row = table_er.add_row()
        row.cells[0].text = entity
        row.cells[1].text = attrs
        row.cells[2].text = types
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9)

    add_formatted_paragraph(doc, '', size=12, space_after=6)

    add_formatted_paragraph(doc, 'Связи между сущностями:', size=12)

    add_bullet_paragraph(doc, 'Category (1) → (N) Product — одна категория содержит множество товаров')
    add_bullet_paragraph(doc, 'auth.User (1) → (N) Order — один пользователь может иметь множество заказов')
    add_bullet_paragraph(doc, 'auth.User (1) → (1) Profile — один пользователь имеет один профиль')
    add_bullet_paragraph(doc, 'Order (1) → (N) OrderItem — один заказ содержит множество позиций')
    add_bullet_paragraph(doc, 'Order (1) → (N) Download — один заказ генерирует множество ссылок для скачивания')
    add_bullet_paragraph(doc, 'Product (1) → (N) OrderItem — один товар может входить в множество позиций заказов')
    add_bullet_paragraph(doc, 'Product (1) → (N) Download — один товар может иметь множество ссылок для скачивания')

    doc.add_page_break()

    # ==================== ПРИЛОЖЕНИЕ Б ====================
    add_heading_styled(doc, 'ПРИЛОЖЕНИЕ Б. Основные фрагменты исходного кода', level=1)

    add_heading_styled(doc, 'Б.1. Класс корзины (cart/cart.py)', level=2)
    add_formatted_paragraph(doc,
        'Класс Cart реализует сессионную корзину покупок. Данные хранятся в сессии пользователя '
        'в формате словаря {product_id: {quantity, price}}. Метод __iter__ обогащает элементы '
        'корзины объектами Product из базы данных и вычисляет итоговую стоимость каждой позиции.', size=12)

    add_heading_styled(doc, 'Б.2. Генерация slug с транслитерацией (products/models.py)', level=2)
    add_formatted_paragraph(doc,
        'Функция generate_unique_slug() обеспечивает автоматическую генерацию уникальных URL-идентификаторов '
        'для товаров и категорий. Выполняет транслитерацию кириллицы, применяет slugify и проверяет '
        'уникальность с добавлением числового суффикса при необходимости.', size=12)

    add_heading_styled(doc, 'Б.3. Механизм скачивания (orders/views.py)', level=2)
    add_formatted_paragraph(doc,
        'Представление download_file реализует безопасное скачивание цифровых товаров. Проверяет '
        'UUID-токен, принадлежность пользователю, срок действия и одноразовость ссылки. Файл отдаётся '
        'через FileResponse как вложение.', size=12)

    add_heading_styled(doc, 'Б.4. Обработка оплаты (payments/views.py)', level=2)
    add_formatted_paragraph(doc,
        'Представление payment_success обрабатывает успешную оплату: меняет статус заказа на «paid», '
        'создаёт записи Download с UUID-токенами и сроком действия 30 дней. Функция payment_webhook '
        'принимает уведомления от платёжной системы и обновляет статус заказа.', size=12)

    add_heading_styled(doc, 'Б.5. Контекстный процессор корзины (cart/context_processors.py)', level=2)
    add_formatted_paragraph(doc,
        'Контекстный процессор cart_processor обеспечивает доступ к объекту корзины и количеству товаров '
        'во всех шаблонах приложения без необходимости явной передачи в каждом представлении.', size=12)

    # Сохранение
    output_path = '/Users/fdevll/diplom/Пояснительная_записка_DigitalStore.docx'
    doc.save(output_path)
    print(f'Документ сохранён: {output_path}')


if __name__ == '__main__':
    create_note()
